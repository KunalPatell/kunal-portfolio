import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import win32com.client, fitz, shutil, os

FONT = 'Calibri'
COLOR_PRIMARY = RGBColor(30, 58, 138)    # Deep Royal Blue (#1E3A8A)
COLOR_ACCENT = RGBColor(2, 132, 199)      # Executive Ocean Blue (#0284C7)
COLOR_LIGHT_BLUE = RGBColor(147, 197, 253) # Light Ocean Blue (#93C5FD)
COLOR_EMERALD = RGBColor(52, 211, 153)    # Emerald (#34D399)
COLOR_SLATE = RGBColor(71, 85, 105)      # Slate Gray (#475569)
COLOR_MUTED = RGBColor(100, 116, 139)     # Muted Gray (#64748B)
COLOR_BODY = RGBColor(30, 41, 59)         # Body Text (#1E293B)
COLOR_WHITE = RGBColor(255, 255, 255)
COLOR_LIGHT_TEXT = RGBColor(241, 245, 249)

HEX_ROYAL_BLUE = '1E3A8A'

def add_hyperlink(paragraph, url, text, font_size=Pt(10.5), bold=False, color_rgb=COLOR_LIGHT_BLUE):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    f = OxmlElement('w:rFonts')
    f.set(qn('w:ascii'), FONT); f.set(qn('w:hAnsi'), FONT)
    rPr.append(f)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
        
    c = OxmlElement('w:color')
    c.set(qn('w:val'), f'{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}')
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=180, bottom=180, left=240, right=240):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def clear_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        node = OxmlElement(f'w:{b}')
        node.set(qn('w:val'), 'none')
        tcBorders.append(node)
    tcPr.append(tcBorders)

projects_data = [
    {
        'title': 'Sevenseed AI Multi-Agent Startup Ecosystem',
        'tech': 'LangGraph • Groq LLaMA 3.3 70B • ChromaDB RAG • FastAPI',
        'url': 'https://sevenseed.onrender.com/',
        'bullets': [
            'Architected a multi-agent SaaS platform orchestrating 9 autonomous AI startups (Comonk, Sevenforce, Breakdown Factor, Rakshak AI, AVP Emart, AVPU, Decode Forest, Trust) using LangGraph, Groq LLaMA 3.3 70B, ChromaDB vector RAG, and FastAPI.',
            'Implemented a self-service Bring-Your-Own-Key (BYOK) key manager in Next.js, enabling zero-cost client workloads across Gemini, OpenAI, and Groq.'
        ]
    },
    {
        'title': 'Comonk AI - Enterprise Career Intelligence Platform',
        'tech': 'RAG Matching • ATS Resume Optimizer • Mock Interview Copilot',
        'url': 'https://comonk-ai.onrender.com/',
        'bullets': [
            'Developed a 32-panel AI career copilot featuring RAG job matching, ATS resume optimizer, voice/video mock interview scoring with PDF reports, and automated application tracking.'
        ]
    },
    {
        'title': 'Sevenforce - Autonomous AI Workforce & Sales CRM',
        'tech': '7-Agent Workforce Dock • Lead Scoring • Bot Automation',
        'url': 'https://sevenseed.onrender.com/sevenforce/',
        'bullets': [
            'Built a 7-agent AI workforce dock (Maya, Sales CRM, automated lead scoring, email automation, meeting bot) for enterprise workflow automation.'
        ]
    },
    {
        'title': 'Breakdown Factor - YOLOv8 Structural Defect Scanner',
        'tech': 'YOLOv8 best.pt • PyTorch • OpenCV • IS-456 BOQ Estimation',
        'url': 'https://sevenseed.onrender.com/breakdown/',
        'bullets': [
            'Fine-tuned a custom YOLOv8 PyTorch model (best.pt) for real-time 10+ category structural defect detection (cracks, pipe leaks, tile damage) and BOQ material cost estimation.'
        ]
    },
    {
        'title': 'Rakshak AI - Citizen Assistant, Police Copilot & Vision Security Suite',
        'tech': 'Automatic FIR (BNS/IPC) • Multilingual Chatbot • PPE Mask Scanner • YOLO',
        'url': 'https://sevenseed.onrender.com/rakshak-ai/',
        'bullets': [
            'Built a 5-in-1 AI platform combining automatic FIR generation with Bharatiya Nyaya Sanhita (BNS/IPC) legal code recommendations, multilingual AI chatbot, cybercrime scam analyzer, and emergency SOS geolocation dispatch.',
            'Integrated Computer Vision workstations for safety mask PPE compliance scanning, sub-second facial attendance verification, and YOLO chair/occupancy detection.'
        ]
    },
    {
        'title': 'AVP Emart - Multi-Store Price & Product Comparison Site',
        'tech': 'E-Commerce Scraper • ML Value Scoring • Streamlit',
        'url': 'https://sevenseed.onrender.com/avp-emart/',
        'bullets': [
            'Engineered an e-commerce price aggregator searching live products across Amazon, Flipkart, Reliance Digital, and Snapdeal with ML value scoring.'
        ]
    }
]

work_data = [
    {
        'role': 'AI-ML Engineer',
        'company': 'Capermint Technology',
        'location': 'Ahmedabad',
        'dates': 'May 2026 - Present',
        'bullets': [
            'Contributed to AI-powered gaming solutions and interactive digital experiences for mobile and web platforms.',
            'Developed and integrated intelligent features, automation workflows, and data-driven solutions to enhance game performance and user engagement.'
        ]
    },
    {
        'role': 'AI Engineer Intern',
        'company': 'Elite Workforces Services',
        'location': 'Ahmedabad',
        'dates': 'Dec 2025 - May 2026',
        'bullets': [
            'Automated business processes using Python, APIs, and n8n, reducing manual effort by ~40%.',
            'Integrated AI services into content and operational workflows; assisted in performance analysis and dashboard insights.'
        ]
    },
    {
        'role': 'AI Intern',
        'company': 'One Percent Media',
        'location': 'Ahmedabad',
        'dates': 'Oct 2025 - Dec 2025',
        'bullets': [
            'Developed and tested automation workflows using Python, n8n, and AI tools.',
            'Integrated AI services into content and operational workflows.'
        ]
    },
    {
        'role': 'AI Automation Engineer',
        'company': 'Sevenseed Technology',
        'location': 'Ahmedabad',
        'dates': 'May 2025 - Nov 2025',
        'bullets': [
            'Designed automation workflows on platforms similar to n8n, Make, Activepieces, and Zapier.',
            'Converted workflow templates into functional automation pipelines and built JSON API-based workflow systems.'
        ]
    }
]

def build_portfolio_resume():
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.28)
        s.bottom_margin = Inches(0.28)
        s.left_margin = Inches(0.35)
        s.right_margin = Inches(0.35)

    style_h1 = [s for s in doc.styles if s.style_id == 'Heading1'][0]
    style_h2 = [s for s in doc.styles if s.style_id == 'Heading2'][0]
    style_h3 = [s for s in doc.styles if s.style_id == 'Heading3'][0]
    style_h4 = [s for s in doc.styles if s.style_id == 'Heading4'][0]
    style_normal = [s for s in doc.styles if s.style_id == 'Normal'][0]
    style_body = [s for s in doc.styles if s.style_id == 'BodyText'][0]
    style_list = [s for s in doc.styles if s.style_id == 'ListParagraph'][0]

    style_h1.font.name = FONT; style_h1.font.size = Pt(28); style_h1.font.bold = True; style_h1.font.color.rgb = COLOR_WHITE
    style_h2.font.name = FONT; style_h2.font.size = Pt(13.5); style_h2.font.bold = True; style_h2.font.color.rgb = COLOR_WHITE
    style_h3.font.name = FONT; style_h3.font.size = Pt(12.5); style_h3.font.bold = True; style_h3.font.color.rgb = COLOR_PRIMARY
    style_h4.font.name = FONT; style_h4.font.size = Pt(11); style_h4.font.bold = True; style_h4.font.italic = True; style_h4.font.color.rgb = COLOR_SLATE

    def add_shaded_h2(title):
        p = doc.add_paragraph(style=style_h2)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.keep_with_next = True
        
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), HEX_ROYAL_BLUE)
        pPr.append(shd)
        
        p.paragraph_format.left_indent = Inches(0.08)
        run_accent = p.add_run('▌ ')
        run_accent.font.name = FONT; run_accent.font.size = Pt(13.5); run_accent.font.bold = True; run_accent.font.color.rgb = COLOR_LIGHT_BLUE
        
        run = p.add_run(title.upper())
        run.font.name = FONT; run.font.size = Pt(13.5); run.font.bold = True; run.font.color.rgb = COLOR_WHITE
        return p

    def add_bullet(text_runs):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.2)
        p.paragraph_format.line_spacing = 1.08
        p.paragraph_format.left_indent = Inches(0.18)
        
        for run_info in text_runs:
            if isinstance(run_info, tuple):
                text, bold, color = run_info
                r = p.add_run(text)
                r.font.name = FONT; r.font.size = Pt(10.5); r.font.bold = bold
                if color: r.font.color.rgb = color
                else: r.font.color.rgb = COLOR_BODY
            else:
                r = p.add_run(run_info)
                r.font.name = FONT; r.font.size = Pt(10.5); r.font.color.rgb = COLOR_BODY
        return p

    # BANNER HEADER
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(7.80)
    set_cell_background(cell, HEX_ROYAL_BLUE)
    set_cell_margins(cell, top=180, bottom=180, left=240, right=240)
    clear_cell_borders(cell)

    p_name = cell.paragraphs[0]
    p_name.style = style_h1
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(1)
    r_name = p_name.add_run('KUNAL PATEL')
    r_name.font.name = FONT; r_name.font.size = Pt(28); r_name.font.bold = True; r_name.font.color.rgb = COLOR_WHITE

    p_title = cell.add_paragraph(style=style_normal)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    r_t = p_title.add_run('AI ENGINEER & AUTOMATION SPECIALIST  |  DATA SCIENTIST')
    r_t.font.name = FONT; r_t.font.size = Pt(12); r_t.font.bold = True; r_t.font.color.rgb = COLOR_LIGHT_BLUE

    p_badge = cell.add_paragraph(style=style_normal)
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_badge.paragraph_format.space_after = Pt(4)
    r_b = p_badge.add_run('⚡ AVAILABLE FOR FULL-TIME AI/ML ROLES & FREELANCE PROJECTS')
    r_b.font.name = FONT; r_b.font.size = Pt(10); r_b.font.bold = True; r_b.font.color.rgb = COLOR_EMERALD

    p_contact = cell.add_paragraph(style=style_normal)
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(1)

    r = p_contact.add_run('📞 8490861586   •   ✉️ ')
    r.font.name = FONT; r.font.size = Pt(10.5); r.font.color.rgb = COLOR_LIGHT_TEXT
    add_hyperlink(p_contact, 'mailto:websitekunal@gmail.com', 'websitekunal@gmail.com', font_size=Pt(10.5), color_rgb=COLOR_LIGHT_BLUE)
    
    r = p_contact.add_run('   •   🔗 ')
    r.font.name = FONT; r.font.size = Pt(10.5); r.font.color.rgb = COLOR_LIGHT_TEXT
    add_hyperlink(p_contact, 'https://www.linkedin.com/in/kunalpatell/', 'LinkedIn', font_size=Pt(10.5), color_rgb=COLOR_LIGHT_BLUE)
    
    r = p_contact.add_run('   •   🌐 ')
    r.font.name = FONT; r.font.size = Pt(10.5); r.font.color.rgb = COLOR_LIGHT_TEXT
    add_hyperlink(p_contact, 'https://kunalpatel-portfolio.vercel.app/', 'Portfolio', font_size=Pt(10.5), color_rgb=COLOR_LIGHT_BLUE)
    
    r = p_contact.add_run('   •   💻 ')
    r.font.name = FONT; r.font.size = Pt(10.5); r.font.color.rgb = COLOR_LIGHT_TEXT
    add_hyperlink(p_contact, 'https://github.com/Kunalptl777', 'GitHub', font_size=Pt(10.5), color_rgb=COLOR_LIGHT_BLUE)

    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(2)
    p_sp.paragraph_format.space_after = Pt(0)

    # 1. ABOUT & EXECUTIVE SUMMARY
    add_shaded_h2('About & Executive Summary')
    p_sum = doc.add_paragraph(style=style_body)
    p_sum.paragraph_format.space_after = Pt(2.5)
    p_sum.paragraph_format.line_spacing = 1.08
    r_sum = p_sum.add_run('AI Engineer with an MSc in Artificial Intelligence & Machine Learning, specializing in shipping end-to-end production AI applications — LLM API gateways (Groq LLaMA 3.3 70B, OpenAI, Gemini), LangGraph autonomous multi-agent systems, custom YOLOv8 Computer Vision models, ChromaDB vector RAG search pipelines, and automated FastAPI backends. Proven track record of architecting 9+ production startup platforms and workflow automation suites.')
    r_sum.font.name = FONT; r_sum.font.size = Pt(10.5); r_sum.font.color.rgb = COLOR_BODY

    # 2. TECHNICAL CAPABILITIES & TOOLKIT
    add_shaded_h2('Technical Capabilities & Toolkit')
    add_bullet([('AI & Multi-Agent Orchestration: ', True, COLOR_PRIMARY), ('LangGraph (Multi-Agent Systems), LangChain, Groq LLaMA 3.3 70B, OpenAI API, Google Gemini, BYOK Architecture, RAG Search, Prompt Engineering.', False, COLOR_BODY)])
    add_bullet([('Computer Vision & Deep Learning: ', True, COLOR_PRIMARY), ('PyTorch, YOLOv8 (Defect Scanner & Occupancy), OpenCV, Image Segmentation, Bounding Box Annotation, Model Fine-Tuning & Evaluation.', False, COLOR_BODY)])
    add_bullet([('Full-Stack Web Stack & Cloud: ', True, COLOR_PRIMARY), ('Python, TypeScript, Next.js, React, Tailwind CSS, FastAPI, Node.js, REST APIs, SQL, Docker, Postman, Git/GitHub, Render, Vercel.', False, COLOR_BODY)])
    add_bullet([('Workflow Automation Engine: ', True, COLOR_PRIMARY), ('n8n, Make, Zapier, JSON API Pipelines, Web Scraping & Data Pipelines, Micro-App Automation.', False, COLOR_BODY)])

    # 3. EDUCATION (MOVED UP)
    add_shaded_h2('Education')
    p_e1 = doc.add_paragraph(style=style_h3)
    p_e1.paragraph_format.space_before = Pt(2.5); p_e1.paragraph_format.space_after = Pt(0.5); p_e1.paragraph_format.keep_with_next = True
    p_e1.paragraph_format.tab_stops.add_tab_stop(Inches(7.80), WD_TAB_ALIGNMENT.RIGHT)
    r = p_e1.add_run('MSc in Artificial Intelligence & Machine Learning'); r.font.name = FONT; r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = COLOR_PRIMARY
    rd = p_e1.add_run('\tAug 2024 - Apr 2026'); rd.font.name = FONT; rd.font.size = Pt(10); rd.font.bold = True; rd.font.color.rgb = COLOR_MUTED

    p_u1 = doc.add_paragraph(style=style_h4)
    p_u1.paragraph_format.space_before = Pt(0); p_u1.paragraph_format.space_after = Pt(1); p_u1.paragraph_format.keep_with_next = True
    r2 = p_u1.add_run('Sardar Patel University, Anand, Gujarat'); r2.font.name = FONT; r2.font.size = Pt(10.5); r2.font.italic = True; r2.font.color.rgb = COLOR_SLATE
    add_bullet([('Focus on AI systems, deep learning, predictive analytics, computer vision, and real-world multi-agent applications.', False, COLOR_BODY)])

    p_e2 = doc.add_paragraph(style=style_h3)
    p_e2.paragraph_format.space_before = Pt(2.5); p_e2.paragraph_format.space_after = Pt(0.5); p_e2.paragraph_format.keep_with_next = True
    p_e2.paragraph_format.tab_stops.add_tab_stop(Inches(7.80), WD_TAB_ALIGNMENT.RIGHT)
    r = p_e2.add_run('Bachelor of Computer Applications (BCA)'); r.font.name = FONT; r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = COLOR_PRIMARY
    rd = p_e2.add_run('\tJun 2020 - Apr 2023'); rd.font.name = FONT; rd.font.size = Pt(10); rd.font.bold = True; rd.font.color.rgb = COLOR_MUTED

    p_u2 = doc.add_paragraph(style=style_h4)
    p_u2.paragraph_format.space_before = Pt(0); p_u2.paragraph_format.space_after = Pt(1); p_u2.paragraph_format.keep_with_next = True
    r2 = p_u2.add_run('Dharmsinh Desai University, Nadiad, Gujarat'); r2.font.name = FONT; r2.font.size = Pt(10.5); r2.font.italic = True; r2.font.color.rgb = COLOR_SLATE
    add_bullet([('Strong foundation in software development, databases, algorithms, and backend technologies.', False, COLOR_BODY)])

    # 4. FEATURED PROJECTS
    add_shaded_h2('Featured Production AI Projects')
    for proj in projects_data:
        p_h = doc.add_paragraph(style=style_h3)
        p_h.paragraph_format.space_before = Pt(2.5)
        p_h.paragraph_format.space_after = Pt(0.5)
        p_h.paragraph_format.keep_with_next = True
        p_h.paragraph_format.tab_stops.add_tab_stop(Inches(7.80), WD_TAB_ALIGNMENT.RIGHT)
        
        r_t = p_h.add_run(proj['title'])
        r_t.font.name = FONT; r_t.font.size = Pt(12); r_t.font.bold = True; r_t.font.color.rgb = COLOR_PRIMARY
        
        p_h.add_run('\t')
        add_hyperlink(p_h, proj['url'], 'Live Project ↗', font_size=Pt(10), bold=True, color_rgb=COLOR_ACCENT)
            
        p_tech = doc.add_paragraph(style=style_normal)
        p_tech.paragraph_format.space_before = Pt(0)
        p_tech.paragraph_format.space_after = Pt(1)
        p_tech.paragraph_format.keep_with_next = True
        r_b = p_tech.add_run(f"TECH STACK: {proj['tech']}")
        r_b.font.name = FONT; r_b.font.size = Pt(9); r_b.font.bold = True; r_b.font.color.rgb = COLOR_SLATE
        
        for btext in proj['bullets']:
            add_bullet([(btext, False, COLOR_BODY)])

    # 5. WORK EXPERIENCE
    p_work = add_shaded_h2('Work Experience')

    for w in work_data:
        p_w = doc.add_paragraph(style=style_h3)
        p_w.paragraph_format.space_before = Pt(2.5)
        p_w.paragraph_format.space_after = Pt(0.5)
        p_w.paragraph_format.keep_with_next = True
        p_w.paragraph_format.tab_stops.add_tab_stop(Inches(7.80), WD_TAB_ALIGNMENT.RIGHT)
        
        r_r = p_w.add_run(w['role'])
        r_r.font.name = FONT; r_r.font.size = Pt(12); r_r.font.bold = True; r_r.font.color.rgb = COLOR_PRIMARY
        
        r_d = p_w.add_run(f"\t{w['dates']}")
        r_d.font.name = FONT; r_d.font.size = Pt(10); r_d.font.bold = True; r_d.font.color.rgb = COLOR_MUTED
        
        p_c = doc.add_paragraph(style=style_h4)
        p_c.paragraph_format.space_before = Pt(0)
        p_c.paragraph_format.space_after = Pt(1)
        p_c.paragraph_format.keep_with_next = True
        r_c = p_c.add_run(f"{w['company']}, {w['location']}")
        r_c.font.name = FONT; r_c.font.size = Pt(10.5); r_c.font.italic = True; r_c.font.color.rgb = COLOR_SLATE
        
        for b in w['bullets']:
            add_bullet([(b, False, COLOR_BODY)])

    temp_docx = r'e:\Project\Portfolio\temp_resume.docx'
    temp_pdf = r'e:\Project\Portfolio\temp_resume.pdf'
    
    doc.save(temp_docx)

    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False

    d = word.Documents.Open(temp_docx)
    d.SaveAs(temp_pdf, FileFormat=17)
    d.Close()
    word.Quit()

    targets = [
        r'e:\Project\Portfolio\public\resume.pdf',
        r'e:\Project\Portfolio\frontend\public\resume.pdf',
        r'e:\Project\Portfolio\public\profile.pdf',
        r'e:\Project\Portfolio\frontend\public\profile.pdf',
        r'e:\Project\Portfolio\public\Kunal_Patell_V9_Executive.pdf',
        r'e:\Project\Portfolio\frontend\public\Kunal_Patell_V9_Executive.pdf'
    ]

    for t in targets:
        shutil.copyfile(temp_pdf, t)
        print(f'Successfully updated: {t}')

    doc_pdf = fitz.open(temp_pdf)
    print(f'New Portfolio Resume PDF page count: {len(doc_pdf)}')

    os.remove(temp_docx)
    print('Updated all Portfolio resume PDF files!')

build_portfolio_resume()
